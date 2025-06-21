from flask import Flask, render_template, request, jsonify
from PyPDF2 import PdfReader
import docx2txt
import io
import base64
from docx import Document
from docx.shared import Pt
from docx.shared import Pt

app = Flask(__name__)

braille_dict = {
    'က': '⡁', 'ခ': '⢈', 'ဂ': '⠛', 'ဃ': '⠟', 'င': '⡈', 'စ': '⡌', 'ဆ': '⡤',
    'ဇ': '⠵', 'ဈ': '⣌', 'ည': '⠷', 'ဋ': '⠳', 'ဌ': '⠻', 'ဍ': '⠾', 'ဎ': '⠿',
    'ဏ': '⡬', 'တ': '⠞', 'ထ': '⠚', 'ဒ': '⠙', 'ဓ': '⠋', 'န': '⠝', 'ပ': '⡖',
    'ဖ': '⠰', 'ဗ': '⢉', 'ဘ': '⠃', 'မ': '⡉', 'ယ': '⠽', 'ရ': '⠗', 'လ': '⠇',
    'ဝ': '⠺', 'သ': '⠹', 'ဟ': '⠓', 'ဠ': '⠸', 'အ': '⠣', 'ဉ': '⠧', 'ဤ': '⠰⠪',
    '၍': '⠯', '၏': '⠕', '၌': '⠦', '၎င်း': '⠬', '၊': '?', '။': '?', 'ာ': '⠁',
    'ါ': '⠎', 'ိ': '⠊', 'ီ': '⠪', 'ု': '⠑', 'ူ': '⠥', 'ေ': '⠱', 'ဲ': '⠡',
    '?': '⠴', 'ံ': '⠉', 'င်္': '⡈⠄⠤', 'ျ': '⠔', 'ဥ': '⠰⠑', 'ဦး': '⠰⠑⠪⠆',
    'ဧ': '⠰⠱', 'ဣ': '⠰⠊', 'ြ': '⠢', '်': '⠄', 'ွ': '⠜', 'ှ': '⠭', '့': '⠂',
    '္': '⠤', 'း': '⠆'
}

def convert_to_braille(text):
    return ''.join(braille_dict.get(c, c) for c in text)

# ✅ Use this to create DOCX with 1.5mm font size (~4.25 pt)
def generate_braille_docx(braille_text):
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(braille_text)

    # Font size 8.5pt ≈ 3mm
    run.font.size = Pt(20)

    # Increase line spacing for realism
    paragraph.paragraph_format.line_spacing = Pt(35)

    return doc

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    text_input = request.form.get("text", "").strip()
    file = request.files.get("file")
    extracted_text = text_input

    if file:
        if file.filename.endswith('.pdf'):
            reader = PdfReader(file)
            extracted_text += "\n" + "".join([page.extract_text() or "" for page in reader.pages])
        elif file.filename.endswith('.docx'):
            extracted_text += "\n" + docx2txt.process(file)
        elif file.filename.endswith('.txt'):
            extracted_text += "\n" + file.read().decode("utf-8")

    braille_text = convert_to_braille(extracted_text)

    # ✅ Generate the DOCX using updated function
    doc = generate_braille_docx(braille_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    encoded_docx = base64.b64encode(buffer.read()).decode("utf-8")

    return jsonify({"braille": braille_text, "docx": encoded_docx})

if __name__ == '__main__':
    app.run(debug=True)
